# -*- coding: utf-8 -*-
"""Acceptance tests for the GOST general-data additions to design documents.

The legacy TZ-14 assertions intentionally remain in ``test_design_docs.py``.
This module exercises only the additive TZ-27 contract and a narrow projection
of the old JSON/XLSX contract that must not change.
"""

from __future__ import annotations

from dataclasses import replace
import hashlib
import json
from pathlib import Path
import re
from typing import Any, Iterable

import pytest

import hvac.design_docs as design_docs_module
from hvac.design_docs import (
    DOCX_OUTPUT_NAME,
    ENGINEER_REVIEW_NOTE,
    JSON_OUTPUT_NAME,
    XLSX_OUTPUT_NAME,
    DesignDocs,
    DesignDocsInputError,
    main,
)


FIXTURES = Path(__file__).resolve().parent / "fixtures"
TERMINAL_RESPONSE = FIXTURES / "hvac-terminal-layout.response.json"
ROUTE_RESPONSE = FIXTURES / "hvac-route-network.response.json"
CHORSU_REQUEST = FIXTURES / "chorsu-l02-terminal-layout.request.json"
CHORSU_TERMINAL_RESPONSE = FIXTURES / "chorsu-l02-terminal-layout.response.json"
CHORSU_ROUTE_RESPONSE = FIXTURES / "chorsu-l02-route-network.response.json"

EXPECTED_SHEET_NAMES = [
    "Спецификация",
    "Воздуховоды",
    "Основные показатели",
    "Ф.1 Чертежи",
    "Ф.2 Документы",
]
MAIN_INDICATOR_HEADERS = (
    "Система",
    "Расход, м³/ч",
    "Давление вентилятора, Па",
    "Терминалов",
    "Обслуживаемая площадь, м²",
)
FORM_1_HEADERS = ("Лист", "Наименование", "Примечание")
FORM_2_HEADERS = ("Обозначение", "Наименование", "Примечание")
FIXED_REFERENCE_DOCUMENTS = (
    "ГОСТ 21.110-2013",
    "ГОСТ 21.602-2016",
)


def _write_json(path: Path, payload: Any) -> Path:
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return path


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _load_workbook(documents: DesignDocs, path: Path) -> Any:
    from openpyxl import load_workbook

    documents.to_xlsx(path)
    return load_workbook(path, data_only=True)


def _values(cells: Iterable[Any]) -> tuple[Any, ...]:
    return tuple(cell.value for cell in cells)


def _find_header_row(sheet: Any, headers: tuple[str, ...]) -> int:
    width = len(headers)
    for row_number, row in enumerate(sheet.iter_rows(), start=1):
        if _values(row[:width]) == headers:
            return row_number
    rendered = [
        _values(row[:width])
        for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row)
    ]
    raise AssertionError(f"Не найдены заголовки {headers!r}; строки: {rendered!r}")


def _table_body(sheet: Any, headers: tuple[str, ...]) -> list[tuple[Any, ...]]:
    header_row = _find_header_row(sheet, headers)
    width = len(headers)
    rows = []
    for row in sheet.iter_rows(min_row=header_row + 1, max_col=width):
        values = _values(row)
        if any(value is not None and value != "" for value in values):
            rows.append(values)
    return rows


def _sheet_text(sheet: Any) -> str:
    return "\n".join(
        str(cell.value)
        for row in sheet.iter_rows()
        for cell in row
        if cell.value is not None
    )


def _normalise_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value)).strip()


def _number_from_text(value: str) -> float:
    match = re.search(r"-?\d+(?:[.,]\d+)?", value.replace(" ", ""))
    if match is None:
        raise AssertionError(f"В ячейке нет числа: {value!r}")
    return float(match.group(0).replace(",", "."))


def test_xlsx_appends_three_gost_sheets_and_preserves_legacy_tables(
    tmp_path: Path,
) -> None:
    documents = DesignDocs.from_files(TERMINAL_RESPONSE, ROUTE_RESPONSE)
    workbook = _load_workbook(documents, tmp_path / "all-sheets.xlsx")

    assert workbook.sheetnames == EXPECTED_SHEET_NAMES

    specification = workbook["Спецификация"]
    ducts = workbook["Воздуховоды"]
    assert _values(specification[5]) == (
        "Поз.",
        "Наименование",
        "Тип, марка",
        "Кол-во",
        "Ед.",
        "Примечание",
    )
    assert _values(ducts[5]) == (
        "Поз.",
        "Сечение",
        "Категория",
        "Участков",
        "Длина, м",
        "Ед.",
        "Примечание",
    )
    assert specification.max_row == 8
    assert ducts.max_row == 18

    assert _find_header_row(
        workbook["Основные показатели"], MAIN_INDICATOR_HEADERS
    )
    assert _find_header_row(workbook["Ф.1 Чертежи"], FORM_1_HEADERS)
    assert _find_header_row(workbook["Ф.2 Документы"], FORM_2_HEADERS)


def test_default_form_1_contains_only_known_outputs_with_honest_placeholders(
    tmp_path: Path,
) -> None:
    documents = DesignDocs.from_files(TERMINAL_RESPONSE, ROUTE_RESPONSE)
    workbook = _load_workbook(documents, tmp_path / "default-form-1.xlsx")
    rows = _table_body(workbook["Ф.1 Чертежи"], FORM_1_HEADERS)

    assert len(rows) == 2
    assert [row[1] for row in rows] == [
        "План этажа",
        "Спецификация оборудования",
    ]
    assert [row[2] for row in rows] == ["дополнить", "дополнить"]
    assert all(row[0] in (None, "", "дополнить") for row in rows)


@pytest.mark.parametrize("container", ["object", "array"])
def test_custom_form_1_accepts_object_and_top_level_array_without_reordering(
    tmp_path: Path,
    container: str,
) -> None:
    sheet_rows = [
        {
            "number": "ОВ-1",
            "name": "План систем вентиляции второго этажа",
            "note": "Изм. 1",
        },
        {
            "number": "ОВ-СО",
            "name": "Спецификация оборудования",
        },
    ]
    payload: Any = {"sheets": sheet_rows} if container == "object" else sheet_rows
    sheets_path = _write_json(tmp_path / f"sheets-{container}.json", payload)

    documents = DesignDocs.from_files(
        TERMINAL_RESPONSE,
        ROUTE_RESPONSE,
        sheets_path,
    )
    workbook = _load_workbook(documents, tmp_path / f"custom-{container}.xlsx")
    rows = _table_body(workbook["Ф.1 Чертежи"], FORM_1_HEADERS)

    assert rows == [
        (
            "ОВ-1",
            "План систем вентиляции второго этажа",
            "Изм. 1",
        ),
        ("ОВ-СО", "Спецификация оборудования", None),
    ]


def test_form_2_does_not_claim_sources_when_velocity_lookup_found_no_value(
    tmp_path: Path,
) -> None:
    documents = DesignDocs.from_files(TERMINAL_RESPONSE, ROUTE_RESPONSE)
    workbook = _load_workbook(documents, tmp_path / "form-2-actual.xlsx")
    sheet = workbook["Ф.2 Документы"]
    rows = _table_body(sheet, FORM_2_HEADERS)
    designations = [str(row[0]).strip() for row in rows if row[0]]

    actual_velocity_documents = {
        entry.source.document
        for entry in design_docs_module.SHNQ_DUCTS.entries
        if entry.key.startswith("velocity.")
    }
    assert actual_velocity_documents
    for document_name in FIXED_REFERENCE_DOCUMENTS:
        assert designations.count(document_name) == 1
    assert actual_velocity_documents.isdisjoint(designations)

    # Legacy TZ-14 warning remains unchanged, while Form 2 is scoped to
    # references actually returned by the lookup.
    assert documents.normative_note == ENGINEER_REVIEW_NOTE
    assert documents.reference_documents_note == ""
    assert ENGINEER_REVIEW_NOTE not in _sheet_text(sheet)


def test_form_2_uses_participating_entry_sources_and_deduplicates_them(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    velocity_entry = next(
        entry
        for entry in design_docs_module.SHNQ_DUCTS.entries
        if entry.key.startswith("velocity.")
    )
    unrelated_entry = next(
        entry
        for entry in design_docs_module.SHNQ_DUCTS.entries
        if entry.key.startswith("velocity.") and entry.key != velocity_entry.key
    )
    used_entry = replace(
        velocity_entry,
        source=replace(
            velocity_entry.source,
            document="ИСПОЛЬЗОВАННЫЙ НОРМАТИВ",
            edition="контрольная редакция",
        ),
        status="unverified",
    )
    unused_entry = replace(
        unrelated_entry,
        source=replace(
            unrelated_entry.source,
            document="НЕИСПОЛЬЗОВАННЫЙ НОРМАТИВ",
            edition="контрольная редакция",
        ),
    )
    catalog = replace(
        design_docs_module.SHNQ_DUCTS,
        document="ОБЩЕЕ ИМЯ КАТАЛОГА",
        entries=(used_entry, unused_entry),
    )
    monkeypatch.setattr(design_docs_module, "SHNQ_DUCTS", catalog)
    monkeypatch.setattr(
        design_docs_module,
        "duct_velocity_limits",
        lambda _kind, _building: (used_entry,),
    )

    documents = design_docs_module.DesignDocs.from_files(
        TERMINAL_RESPONSE,
        ROUTE_RESPONSE,
    )
    workbook = _load_workbook(documents, tmp_path / "form-2-traced-source.xlsx")
    rows = _table_body(workbook["Ф.2 Документы"], FORM_2_HEADERS)
    designations = [str(row[0]).strip() for row in rows if row[0]]

    assert designations.count("ИСПОЛЬЗОВАННЫЙ НОРМАТИВ") == 1
    assert "НЕИСПОЛЬЗОВАННЫЙ НОРМАТИВ" not in designations
    assert "ОБЩЕЕ ИМЯ КАТАЛОГА" not in designations
    assert ENGINEER_REVIEW_NOTE in _sheet_text(workbook["Ф.2 Документы"])


def test_unused_unverified_velocity_source_does_not_warn_in_form_2(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    velocity_entries = [
        entry
        for entry in design_docs_module.SHNQ_DUCTS.entries
        if entry.key.startswith("velocity.")
    ]
    first, second = velocity_entries[:2]
    used_verified = replace(
        first,
        source=replace(first.source, document="ПРИМЕНЁННЫЙ VERIFIED НОРМАТИВ"),
        status="verified",
        verified_by="Инженер ОВ",
        verified_at="2026-08-23T10:00:00Z",
    )
    unused_unverified = replace(
        second,
        source=replace(second.source, document="НЕПРИМЕНЁННЫЙ UNVERIFIED НОРМАТИВ"),
        status="unverified",
        verified_by=None,
        verified_at=None,
    )
    catalog = replace(
        design_docs_module.SHNQ_DUCTS,
        entries=(used_verified, unused_unverified),
    )
    monkeypatch.setattr(design_docs_module, "SHNQ_DUCTS", catalog)
    monkeypatch.setattr(
        design_docs_module,
        "duct_velocity_limits",
        lambda _kind, _building: (used_verified,),
    )

    documents = design_docs_module.DesignDocs.from_files(
        TERMINAL_RESPONSE,
        ROUTE_RESPONSE,
    )
    workbook = _load_workbook(documents, tmp_path / "form-2-used-only.xlsx")
    sheet = workbook["Ф.2 Документы"]
    rendered = _sheet_text(sheet)

    assert "ПРИМЕНЁННЫЙ VERIFIED НОРМАТИВ" in rendered
    assert "НЕПРИМЕНЁННЫЙ UNVERIFIED НОРМАТИВ" not in rendered
    assert documents.normative_note == ENGINEER_REVIEW_NOTE
    assert documents.reference_documents_note == ""
    assert ENGINEER_REVIEW_NOTE not in rendered


def test_fully_verified_velocity_sources_remove_form_2_warning(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    verified_entries = tuple(
        (
            replace(
                entry,
                status="verified",
                verified_by="Инженер ОВ",
                verified_at="2026-08-23T10:00:00Z",
            )
            if entry.key.startswith("velocity.")
            else entry
        )
        for entry in design_docs_module.SHNQ_DUCTS.entries
    )
    verified_velocity_entries = tuple(
        entry for entry in verified_entries if entry.key.startswith("velocity.")
    )
    catalog = replace(
        design_docs_module.SHNQ_DUCTS,
        entries=verified_entries,
    )
    monkeypatch.setattr(design_docs_module, "SHNQ_DUCTS", catalog)
    monkeypatch.setattr(
        design_docs_module,
        "duct_velocity_limits",
        lambda _kind, _building: verified_velocity_entries,
    )

    documents = design_docs_module.DesignDocs.from_files(
        TERMINAL_RESPONSE,
        ROUTE_RESPONSE,
    )
    workbook = _load_workbook(documents, tmp_path / "form-2-verified.xlsx")
    sheet = workbook["Ф.2 Документы"]
    designations = [
        str(row[0]).strip()
        for row in _table_body(sheet, FORM_2_HEADERS)
        if row[0]
    ]

    assert documents.normative_note == ""
    assert ENGINEER_REVIEW_NOTE not in _sheet_text(sheet)
    assert designations.count(design_docs_module.SHNQ_DUCTS.document) == 1


def test_legacy_two_response_call_keeps_tz14_json_and_marks_missing_area(
    tmp_path: Path,
) -> None:
    documents = DesignDocs.from_files(TERMINAL_RESPONSE, ROUTE_RESPONSE)
    payload = documents.as_dict()

    assert payload["schemaVersion"] == "1.0"
    assert payload["designNote"]["systems"] == [
        {
            "systemId": "SA-01",
            "airflowM3h": 1080.0,
            "fanPressurePa": pytest.approx(125.878512153944),
            "dictatingBranch": {
                "segmentIds": list(documents.systems[0].critical_path_segment_ids),
                "segmentCount": 12,
                "lengthM": 25.2,
            },
            "terminalCount": 6,
        }
    ]
    assert payload["specification"]["standard"] == "ГОСТ 21.110-2013"
    assert payload["ductSchedule"]["totalScheduledLengthM"] == 40.8

    workbook = _load_workbook(documents, tmp_path / "legacy-no-request.xlsx")
    rows = _table_body(workbook["Основные показатели"], MAIN_INDICATOR_HEADERS)
    sa_01 = next(row for row in rows if row[0] == "SA-01")
    assert sa_01[:4] == (
        "SA-01",
        1080.0,
        pytest.approx(125.878512153944),
        6,
    )
    assert sa_01[4] == "дополнить"


def test_route_request_populates_area_from_authoritative_placement_refs(
    tmp_path: Path,
) -> None:
    terminal = _read_json(TERMINAL_RESPONSE)
    route = _read_json(ROUTE_RESPONSE)
    request = {
        "schemaVersion": route["schemaVersion"],
        "profileVersion": route["profileVersion"],
        "kind": "hvac-route-network-request",
        "requestId": route["requestId"],
        "sourceEvidence": route["sourceEvidence"],
        "spaces": [
            {
                "spaceId": "ROOM-A",
                "geometry": {
                    "kind": "rectangular-local",
                    "widthMm": 3000,
                    "lengthMm": 4000,
                },
            },
            {
                "spaceId": "ROOM-B",
                "geometry": {
                    "kind": "rectangular-local",
                    "widthMm": 3000,
                    "lengthMm": 4000,
                },
            },
            {
                "spaceId": "ROOM-C",
                "geometry": {
                    "kind": "rectangular-local",
                    "widthMm": 4000,
                    "lengthMm": 4000,
                },
            },
        ],
        "placements": [
            {"placementId": "terminal-a-1", "spaceId": "ROOM-A"},
            {"placementId": "terminal-a-2", "spaceId": "ROOM-A"},
            {"placementId": "terminal-b-1", "spaceId": "ROOM-B"},
            {"placementId": "terminal-b-2", "spaceId": "ROOM-B"},
            {"placementId": "terminal-c-1", "spaceId": "ROOM-C"},
            {"placementId": "terminal-c-2", "spaceId": "ROOM-C"},
        ],
        "systems": [
            {
                "systemId": "SA-01",
                "placementRefs": [
                    "terminal-a-1",
                    "terminal-a-2",
                    "terminal-b-1",
                    "terminal-b-2",
                    "terminal-c-1",
                    "terminal-c-2",
                ],
            }
        ],
        "sizing": {"trunkVelocityMs": 7, "branchVelocityMs": 4},
        "provenance": {"modelMutationRequested": False},
    }
    request_path = _write_json(tmp_path / "route-request.json", request)
    route["requestSha256"] = hashlib.sha256(request_path.read_bytes()).hexdigest().upper()
    route_path = _write_json(tmp_path / "route-response.json", route)
    terminal["sourceEvidence"] = request["sourceEvidence"]
    terminal_path = _write_json(tmp_path / "terminal-response.json", terminal)

    documents = DesignDocs.from_files(
        terminal_path,
        route_path,
        request_path=request_path,
    )
    indicator = documents.main_indicators[0]

    assert indicator.system_id == "SA-01"
    assert indicator.terminal_count == 6
    assert indicator.served_space_ids == ("ROOM-A", "ROOM-B", "ROOM-C")
    assert indicator.served_area_m2 == pytest.approx(40.0)


def test_tampered_request_geometry_is_rejected_by_source_hash(
    tmp_path: Path,
) -> None:
    request = _read_json(CHORSU_REQUEST)
    request["spaces"][0]["geometry"]["outerBoundaryMm"][0][0] += 1000
    tampered = _write_json(tmp_path / "tampered-request.json", request)

    with pytest.raises(DesignDocsInputError, match="requestSha256"):
        DesignDocs.from_files(
            CHORSU_TERMINAL_RESPONSE,
            CHORSU_ROUTE_RESPONSE,
            request_path=tampered,
        )


def test_request_rejects_other_response_from_different_source_model(
    tmp_path: Path,
) -> None:
    route = _read_json(CHORSU_ROUTE_RESPONSE)
    route["sourceEvidence"]["activeDocumentFingerprint"] = "A" * 64
    unrelated_route = _write_json(tmp_path / "unrelated-route.json", route)

    with pytest.raises(
        DesignDocsInputError,
        match=r"route\.sourceEvidence\.activeDocumentFingerprint",
    ):
        DesignDocs.from_files(
            CHORSU_TERMINAL_RESPONSE,
            unrelated_route,
            request_path=CHORSU_REQUEST,
        )


def test_chorsu_request_populates_system_areas_and_docx_main_indicators(
    tmp_path: Path,
) -> None:
    from docx import Document

    documents = DesignDocs.from_files(
        CHORSU_TERMINAL_RESPONSE,
        CHORSU_ROUTE_RESPONSE,
        request_path=CHORSU_REQUEST,
    )
    workbook = _load_workbook(documents, tmp_path / "chorsu.xlsx")
    rows = _table_body(workbook["Основные показатели"], MAIN_INDICATOR_HEADERS)
    by_system = {str(row[0]): row for row in rows if row[0]}

    assert set(by_system) == {"EA-CHORSU-L02", "SA-CHORSU-L02"}
    assert by_system["EA-CHORSU-L02"][1:4] == (
        2850.0,
        pytest.approx(64.37882890346316),
        28,
    )
    assert by_system["SA-CHORSU-L02"][1:4] == (
        4050.0,
        pytest.approx(71.5975189890657),
        35,
    )
    assert float(by_system["EA-CHORSU-L02"][4]) == pytest.approx(
        156.34616008,
        abs=0.01,
    )
    assert float(by_system["SA-CHORSU-L02"][4]) == pytest.approx(
        776.103316575,
        abs=0.01,
    )

    output = tmp_path / "chorsu.docx"
    documents.to_docx(output)
    document = Document(output)
    paragraphs = [_normalise_text(paragraph.text) for paragraph in document.paragraphs]
    assert "Вентиляция: принятые решения" in paragraphs
    assert any("Основные показатели" in paragraph for paragraph in paragraphs)

    main_table = next(
        table
        for table in document.tables
        if tuple(_normalise_text(cell.text) for cell in table.rows[0].cells)
        == MAIN_INDICATOR_HEADERS
    )
    docx_rows = {
        _normalise_text(row.cells[0].text): [
            _normalise_text(cell.text) for cell in row.cells
        ]
        for row in main_table.rows[1:]
    }
    assert set(docx_rows) == {"EA-CHORSU-L02", "SA-CHORSU-L02"}
    assert _number_from_text(docx_rows["EA-CHORSU-L02"][4]) == pytest.approx(
        156.34616008,
        abs=0.01,
    )
    assert _number_from_text(docx_rows["SA-CHORSU-L02"][4]) == pytest.approx(
        776.103316575,
        abs=0.01,
    )


def test_cli_accepts_sheets_and_request_without_changing_output_names(
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
) -> None:
    sheets_path = _write_json(
        tmp_path / "cli-sheets.json",
        {
            "sheets": [
                {
                    "number": "ОВ-01",
                    "name": "План вентиляции",
                    "note": "Проверено",
                }
            ]
        },
    )
    output_directory = tmp_path / "cli-result"

    result = main(
        [
            "--terminal",
            str(CHORSU_TERMINAL_RESPONSE),
            "--route",
            str(CHORSU_ROUTE_RESPONSE),
            "--sheets",
            str(sheets_path),
            "--request",
            str(CHORSU_REQUEST),
            "--out",
            str(output_directory),
        ]
    )

    assert result == 0
    captured = capsys.readouterr()
    assert captured.err == ""
    output_names = (JSON_OUTPUT_NAME, XLSX_OUTPUT_NAME, DOCX_OUTPUT_NAME)
    assert sorted(path.name for path in output_directory.iterdir()) == sorted(output_names)
    assert all(name in captured.out for name in output_names)

    from openpyxl import load_workbook

    workbook = load_workbook(output_directory / XLSX_OUTPUT_NAME, data_only=True)
    assert workbook.sheetnames == EXPECTED_SHEET_NAMES
    assert _table_body(workbook["Ф.1 Чертежи"], FORM_1_HEADERS) == [
        ("ОВ-01", "План вентиляции", "Проверено")
    ]
    indicator_rows = _table_body(
        workbook["Основные показатели"], MAIN_INDICATOR_HEADERS
    )
    assert all(row[4] != "дополнить" for row in indicator_rows if row[0])
