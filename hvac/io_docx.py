# -*- coding: utf-8 -*-
"""Экспорт пояснительной записки в DOCX (Microsoft Word).

Редактируемый аналог PDF-записки (io_pdf.py): те же разделы и данные,
но в формате, который дорабатывают вручную под требования экспертизы.
Использует python-docx (pip install python-docx).

Содержимое таблиц строится в hvac/report_sections.py (общее с PDF),
подписи нормативных документов — по профилю норм проекта
(hvac/catalogs/norm_profiles.py: УзР ШНҚ/КМК либо РФ СП, по
params.thermal_norm). Номера разделов сквозные — только по реально
напечатанным разделам. Энергопаспорт перед печатью актуализируется
пересчётом по текущим помещениям (hvac.energy.refresh_passport).

Разделы (include_sections): 'cover', 'inputs', 'constructions',
'heat_loss', 'heat_gain', 'ventilation', 'dhw', 'equipment', 'smoke',
'ducts', 'pipes', 'energy', 'condensation'. None — все.
"""

from __future__ import annotations

from datetime import datetime
from typing import List, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from hvac.project import HVACProject

ALL_SECTIONS = [
    "cover", "inputs", "constructions", "heat_loss", "heat_gain",
    "ventilation", "dhw", "equipment", "smoke", "ducts", "pipes",
    "energy", "condensation",
]


def export_to_docx(project: "HVACProject", path: str,
                   include_sections: Optional[List[str]] = None) -> None:
    """Создаёт DOCX-отчёт «Пояснительная записка» по проекту.

    include_sections : список разделов (None = все, см. ALL_SECTIONS).
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("Не установлен python-docx. "
                           "Выполните: pip install python-docx")

    from hvac import __version__
    from hvac import report_sections as rs
    from hvac.catalogs.norm_profiles import get_norm_profile

    if include_sections is None:
        include_sections = list(ALL_SECTIONS)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    p = project.params
    prof = get_norm_profile(p)
    num = rs.SectionNumberer()

    def add_para(text: str, *, bold: bool = False, center: bool = False,
                 size: Optional[int] = None):
        par = doc.add_paragraph()
        if center:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        return par

    def add_table(rows: Optional[List[List]], header_row: bool = True):
        """Таблица с единым стилем; первая строка — шапка."""
        if not rows:
            return
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = t.cell(i, j)
                cell.text = str(val)
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.font.size = Pt(9)
                        if header_row and i == 0:
                            run.bold = True
        doc.add_paragraph()

    def add_kv(rows: List[List]):
        add_table([["Параметр", "Значение"]] + [list(r) for r in rows])

    # Актуализация энергопаспорта ДО титульного листа: суммы на титуле и
    # раздел энергоэффективности должны сходиться с текущими помещениями.
    if "energy" in include_sections and project.energy_passport is not None:
        from hvac.energy import refresh_passport
        refresh_passport(project)

    # =================== ТИТУЛЬНЫЙ ЛИСТ ===================
    if "cover" in include_sections:
        for _ in range(6):
            doc.add_paragraph()
        add_para("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", bold=True, center=True, size=22)
        add_para("Раздел «Отопление, вентиляция и кондиционирование воздуха»",
                 bold=True, center=True, size=14)
        doc.add_paragraph()
        add_para(f"Объект: {p.project_name or 'Объект'}", bold=True)
        add_para(f"Город: {p.city}", bold=True)
        doc.add_paragraph()

        n_spaces = len(project.spaces)
        total_area = sum(s.area_m2 for s in project.spaces)
        total_q = sum(s.heat_loss_w for s in project.spaces) / 1000.0
        total_qg = sum(s.heat_gain_w for s in project.spaces) / 1000.0
        sum_supply = sum(s.supply_m3h for s in project.spaces)
        sum_exhaust = sum(s.exhaust_m3h + s.hood_m3h for s in project.spaces)
        add_para(f"Помещений: {n_spaces}, общая площадь "
                 f"{rs.fmt(total_area)} м²")
        add_para(f"Σ Q отопления: {total_q:.1f} кВт")
        add_para(f"Σ Q охлаждения: {total_qg:.1f} кВт")
        if sum_supply > 0 or sum_exhaust > 0:
            add_para(f"Σ приток: {rs.fmt(sum_supply)} м³/ч, "
                     f"Σ вытяжка: {rs.fmt(sum_exhaust)} м³/ч")
        doc.add_paragraph()
        add_para(f"Расчёт выполнен: {datetime.now().strftime('%d.%m.%Y')}")
        add_para(f"ПО: HVAC Calculator v{__version__}")
        add_para(f"{prof['title']}: {prof['methods_line']}")
        add_para(f"Методика расчёта нагрузок: {p.methodology}")
        doc.add_page_break()

    # =================== ИСХОДНЫЕ ДАННЫЕ ===================
    if "inputs" in include_sections:
        doc.add_heading(num.title("Исходные данные"), level=1)
        add_para(f"Климатические параметры приняты по {prof['climate']} для "
                 f"города {p.city}. Внутренние параметры микроклимата — по "
                 f"{prof['indoor']} для соответствующих типов помещений.")
        add_kv([
            ["Город", p.city],
            ["Расчётная зимняя tн, °C (обесп. 0.92)", f"{p.t_out_heating}"],
            ["Расчётная летняя tн, °C (обесп. 0.95)", f"{p.t_out_cooling}"],
            ["Суточная амплитуда летом, K", f"{p.daily_amplitude}"],
            ["ГСОП (tв=+20°C, ≤8°C), °C·сут", f"{p.gsop_18:.0f}"],
            ["Солнечная радиация на вертик., Вт/м²",
             f"{p.solar_intensity_w_m2}"],
            ["Влагосодержание нар. лето, г/кг", f"{p.w_out_summer_g_kg}"],
            ["Влагосодержание вн. лето, г/кг", f"{p.w_in_summer_g_kg}"],
            ["Методика расчёта", p.methodology],
            ["Коэф. поправки инфильтрации", f"{p.inf_correction_k}"],
            ["Запас на отопление", f"{p.safety_margin_heating}"],
            ["Запас на охлаждение", f"{p.safety_margin_cooling}"],
        ])

    # =================== КОНСТРУКЦИИ ===================
    if "constructions" in include_sections and project.constructions:
        doc.add_heading(num.title("Каталог конструкций"), level=1)
        add_para("Каталог U-значений и SHGC ограждающих конструкций. "
                 "Сформирован автоматически по типам элементов из Revit. "
                 f"Требуемые сопротивления R₀ — по {prof['heat_loss_short']}.")
        rows = [["Категория", "Семейство", "Тип", "Δ, мм", "U", "SHGC"]]
        for key, c in sorted(project.constructions.items(),
                             key=lambda kv: (kv[1].category, kv[1].family)):
            rows.append([
                c.category, c.family, c.type_name[:35],
                f"{c.thickness_mm:.0f}",
                f"{c.u_value:.2f}" if c.u_value > 0 else "—",
                f"{c.shgc:.2f}" if c.shgc > 0 else "—",
            ])
        add_table(rows)

    # =================== ТЕПЛОПОТЕРИ ===================
    if "heat_loss" in include_sections and project.spaces:
        doc.add_heading(
            num.title(f"Теплопотери ({prof['heat_loss_short']})"), level=1)
        loaded = [s for s in project.spaces if s.heat_loss_w > 0]
        total = sum(s.heat_loss_w for s in loaded)
        total_area = sum(s.area_m2 for s in loaded)
        add_para(f"Расчёт теплопотерь выполнен в соответствии с "
                 f"{prof['heat_loss']} с разбивкой по статьям (ограждения, "
                 f"инфильтрация, внутренние перегородки). "
                 f"Σ = {total / 1000:.1f} кВт "
                 f"({total / max(total_area, 1):.1f} Вт/м² удельно).")
        add_table(rs.heat_loss_level_rows(project))
        room_rows = rs.heat_loss_room_rows(project)
        if room_rows:
            add_para("Теплопотери по помещениям:", bold=True)
            add_table(room_rows)
        elif loaded:
            add_para(f"По-помещенная таблица опущена ({len(loaded)} помещений "
                     f"> {rs.ROOM_TABLE_LIMIT}); полные данные — в "
                     "Excel-экспорте.")

    # =================== ТЕПЛОПОСТУПЛЕНИЯ ===================
    if "heat_gain" in include_sections and project.spaces:
        doc.add_heading(
            num.title(f"Теплопоступления ({prof['heat_gain']})"), level=1)
        loaded = [s for s in project.spaces if s.heat_gain_w > 0]
        total = sum(s.heat_gain_w for s in loaded)
        total_sens = sum(s.heat_gain_sensible_w for s in loaded)
        total_lat = sum(s.heat_gain_latent_w for s in loaded)
        add_para(f"Расчёт по {prof['heat_gain']} с разделением на явную "
                 f"(sensible) и скрытую (latent) теплоту: солнечная радиация, "
                 f"люди, освещение, оборудование, инфильтрация. "
                 f"Σ = {total / 1000:.1f} кВт, из них явная "
                 f"{total_sens / 1000:.1f}, скрытая {total_lat / 1000:.1f}.")
        add_table(rs.heat_gain_level_rows(project))
        room_rows = rs.heat_gain_room_rows(project)
        if room_rows:
            add_para("Теплопоступления по помещениям:", bold=True)
            add_table(room_rows)

    # =================== ВЕНТИЛЯЦИЯ ===================
    if "ventilation" in include_sections and project.spaces:
        doc.add_heading(
            num.title(f"Вентиляция ({prof['ventilation']})"), level=1)
        add_para(f"Воздухообмены приняты по {prof['ventilation']} и "
                 "технологическим требованиям: по расчётному числу людей, "
                 "кратностям и нормативам на санитарные приборы.")
        add_para(rs.air_balance_note(project))
        air_rows = rs.air_exchange_room_rows(project)
        if air_rows:
            add_para("Таблица воздухообменов по помещениям:", bold=True)
            add_table(air_rows)
        sys_rows = rs.vent_system_summary_rows(project)
        if sys_rows:
            add_para("Сводка по вентиляционным системам:", bold=True)
            add_table(sys_rows)

    # =================== ГВС ===================
    if "dhw" in include_sections and project.dhw_systems:
        doc.add_heading(
            num.title(f"Горячее водоснабжение ({prof['dhw_short']})"), level=1)
        add_para(f"Нормативный документ: {prof['dhw']}. Удельные нормы "
                 "расхода горячей воды на потребителя — по Прил. А "
                 "СП 30.13330 (гармонизированы со СНиП 2.04.01-85*). "
                 "Учтены потери на циркуляцию и КПД нагревателей.")
        rows = [["Система", "N", "V сут, м³", "V час, м³/ч",
                 "Q пик, кВт", "Q нагр., кВт", "V бака, м³"]]
        total_v = total_q = 0.0
        for name, sys in sorted(project.dhw_systems.items()):
            rows.append([
                name, str(sys.n_consumers),
                f"{sys.v_daily_total_m3:.1f}",
                f"{sys.v_hourly_max_m3:.2f}",
                f"{sys.q_peak_w / 1000:.1f}",
                f"{sys.q_heater_size_w / 1000:.1f}",
                f"{sys.storage_recommended_m3:.2f}",
            ])
            total_v += sys.v_daily_total_m3
            total_q += sys.q_heater_size_w
        rows.append(["ИТОГО", "", f"{total_v:.1f}", "", "",
                     f"{total_q / 1000:.1f}", ""])
        add_table(rows)

    # =================== СИСТЕМЫ ОБОРУДОВАНИЯ ===================
    has_equipment = bool(project.ventilation_systems
                         or project.heating_systems
                         or project.cooling_systems or project.ahu_loads)
    if "equipment" in include_sections and has_equipment:
        doc.add_heading(num.title("Системы оборудования"), level=1)
        if project.ventilation_systems:
            doc.add_heading("Приточные/вытяжные установки", level=2)
            rows = [["Имя", "Тип", "Рекуп.", "η зима",
                     "t под. зима", "t под. лето"]]
            for name, v in sorted(project.ventilation_systems.items()):
                rows.append([
                    name, v.system_type,
                    "Да" if v.has_recovery else "Нет",
                    f"{v.recovery_efficiency_winter:.2f}",
                    f"{v.t_supply_winter}°C", f"{v.t_supply_summer}°C",
                ])
            add_table(rows)
        if project.heating_systems:
            doc.add_heading("Источники тепла", level=2)
            rows = [["Имя", "Тип", "t подачи", "t обр.", "Топливо", "η"]]
            for name, h in sorted(project.heating_systems.items()):
                rows.append([name, h.system_type, f"{h.t_supply}°C",
                             f"{h.t_return}°C", h.fuel, f"{h.efficiency:.2f}"])
            add_table(rows)
        if project.cooling_systems:
            doc.add_heading("Источники холода", level=2)
            rows = [["Имя", "Тип", "t подачи", "t обр.",
                     "Хладагент", "EER/COP"]]
            for name, cs in sorted(project.cooling_systems.items()):
                rows.append([name, cs.system_type, f"{cs.t_supply}°C",
                             f"{cs.t_return}°C", cs.refrigerant,
                             f"{cs.cop:.2f}"])
            add_table(rows)
        if project.ahu_loads:
            doc.add_heading("Нагрузки от приточных установок", level=2)
            rows = [["AHU", "Supply, м³/ч", "Q калориф., кВт",
                     "Q охл. явн., кВт", "Q охл. скр., кВт"]]
            total_h = total_cs = total_cl = 0.0
            for name, d in sorted(project.ahu_loads.items()):
                rows.append([
                    name, rs.fmt(d['supply_m3h']),
                    f"{d['q_heater_w'] / 1000:.1f}",
                    f"{d['q_cooler_sens_w'] / 1000:.1f}",
                    f"{d['q_cooler_lat_w'] / 1000:.1f}",
                ])
                total_h += d["q_heater_w"]
                total_cs += d["q_cooler_sens_w"]
                total_cl += d["q_cooler_lat_w"]
            rows.append(["ИТОГО", "", f"{total_h / 1000:.1f}",
                         f"{total_cs / 1000:.1f}", f"{total_cl / 1000:.1f}"])
            add_table(rows)

    # ============= ПРОТИВОДЫМНАЯ ЗАЩИТА (дымоудаление + подпор) =============
    if "smoke" in include_sections and project.smoke_systems:
        from hvac.catalogs.smoke_norms import get_smoke_norm
        from hvac.smoke_report import build_smoke_explanations
        norm = get_smoke_norm(getattr(project.params, "smoke_norm", "SP7_RU"))
        sec_n = num.title(
            "Противодымная защита: дымоудаление и подпор воздуха")
        doc.add_heading(sec_n, level=1)
        sec_prefix = sec_n.split(".", 1)[0]
        add_para(f"Действующий норматив: {norm.reference}")
        add_para("Расход дымоудаления и подпора воздуха определён по формулам "
                 "норматива. Ниже по каждой системе приведён вывод расчёта с "
                 "исходными данными, подстановкой и ссылкой на пункт норматива. "
                 "Оборудование противодымной защиты включается автоматически "
                 "при пожаре (ШНК 2.01.02-04, п. 5.7).")

        # Пересчёт расходов (заполняет L_smoke_m3h/L_makeup_m3h по системам)
        explanations = build_smoke_explanations(project)

        smoke = {k: v for k, v in project.smoke_systems.items()
                 if v.system_type == "smoke_removal"}
        press = {k: v for k, v in project.smoke_systems.items()
                 if v.system_type == "air_supply"}

        if smoke:
            doc.add_heading(
                f"{sec_prefix}.1. Системы дымоудаления (СДУ) — сводка",
                level=2)
            rows = [["Имя", "Назначение", "Метод", "Площадь, м²", "Зон",
                     "L сист., м³/ч", "L комп., м³/ч", "Огнест."]]
            for name, ss in sorted(smoke.items()):
                rows.append([
                    name, ss.purpose, ss.calc_method,
                    f"{ss.served_area_m2:.0f}", str(ss.n_zones),
                    rs.fmt(ss.L_smoke_m3h), rs.fmt(ss.L_makeup_m3h),
                    ss.fire_rating,
                ])
            add_table(rows)
        if press:
            doc.add_heading(
                f"{sec_prefix}.2. Системы подпора воздуха (СПВ) — сводка",
                level=2)
            rows = [["Имя", "Назначение", "L, м³/ч", "Давление, Па"]]
            for name, ps in sorted(press.items()):
                rows.append([
                    name, ps.purpose, rs.fmt(ps.L_smoke_m3h),
                    f"{ps.pressure_pa:.0f}",
                ])
            add_table(rows)

        if explanations:
            doc.add_heading(
                f"{sec_prefix}.3. Расчёт с пояснениями (по системам)", level=2)
            for ex in explanations:
                doc.add_heading(
                    f"{ex['kind']} «{ex['name']}» — {ex['method_title']}",
                    level=3)
                add_para(f"Формула:  {ex['formula']}", bold=True)
                add_para(f"Норматив: {ex['ref']}")
                add_table([["Исходные данные", "Значение"]]
                          + [list(pair) for pair in ex["inputs"]])
                add_para(f"Подстановка:  {ex['substitution']}")
                add_table([["Результат", "Значение"]]
                          + [list(pair) for pair in ex["results"]])
                if ex.get("checks"):
                    add_para("Требования и проверки:", bold=True)
                    for c in ex["checks"]:
                        add_para("• " + c)
                if ex.get("note"):
                    add_para("Примечание: " + ex["note"])

    # =================== ВОЗДУХОВОДЫ ===================
    if "ducts" in include_sections and project.duct_networks:
        doc.add_heading(
            num.title("Подбор воздуховодов (упрощённая аэродинамика)"),
            level=1)
        add_para("Подбор по рекомендованным скоростям, потери давления — "
                 "формула Дарси-Вейсбаха.")
        for sys_name, net in sorted(project.duct_networks.items()):
            doc.add_heading(
                f"{sys_name}: Σ Q = {rs.fmt(net.total_flow_m3h)} м³/ч, "
                f"Δp ≈ {net.total_pressure_loss_pa:.0f} Па", level=2)
            rows = [["Участок", "Q, м³/ч", "Размер", "v, м/с", "Δp, Па"]]
            for sec in net.sections:
                size_str = (f"Ø{int(sec.diameter_mm)}"
                            if sec.shape == "round"
                            else f"{int(sec.width_mm)}×{int(sec.height_mm)}")
                rows.append([
                    sec.id[-30:], rs.fmt(sec.flow_m3h), size_str,
                    f"{sec.velocity_m_s:.1f}",
                    f"{sec.pressure_loss_total_pa:.0f}",
                ])
            add_table(rows)

    # =================== ТРУБЫ ОТОПЛЕНИЯ ===================
    if "pipes" in include_sections and project.pipe_networks:
        doc.add_heading(
            num.title("Подбор труб отопления (гидравлический расчёт)"),
            level=1)
        from hvac.pipe_sizing import WATER_DENSITY_70C
        add_para("Подбор по рекомендованным скоростям, потери давления — "
                 "формула Альтшуля.")
        for sys_name, pnet in sorted(project.pipe_networks.items()):
            pump_m = pnet.total_pressure_loss_pa / (WATER_DENSITY_70C * 9.81)
            doc.add_heading(
                f"{sys_name}: Σ Q = {pnet.total_heat_load_w / 1000:.1f} кВт, "
                f"Σ G = {pnet.total_flow_kg_h:.0f} кг/ч, "
                f"Δp = {pnet.total_pressure_loss_pa / 1000:.1f} кПа, "
                f"напор ≈ {pump_m:.1f} м", level=2)
            rows = [["Участок", "Q, Вт", "DN", "v, м/с", "Δp, Па"]]
            for psec in pnet.sections:
                rows.append([
                    psec.id[-30:], rs.fmt(psec.heat_load_w),
                    f"{int(psec.dn_mm)}", f"{psec.velocity_m_s:.2f}",
                    f"{psec.pressure_loss_total_pa:.0f}",
                ])
            add_table(rows)

    # =================== ЭНЕРГОЭФФЕКТИВНОСТЬ ===================
    if "energy" in include_sections and project.energy_passport:
        ep = project.energy_passport
        data = rs.energy_section_data(ep, prof)
        doc.add_heading(num.title(data["heading"]), level=1)
        add_para(data["intro"])
        add_kv(data["kv_main"])
        add_para(data["verdict"], bold=True, center=True, size=12)
        add_para("Годовое энергопотребление (справочно):", bold=True)
        add_kv(data["kv_annual"])
        for note in data["notes"]:
            add_para("Примечание: " + note)

    # =================== ТОЧКА РОСЫ ===================
    if "condensation" in include_sections and project.condensation_results:
        doc.add_heading(
            num.title("Проверка ограждений на конденсацию "
                      f"({prof['condensation']})"), level=1)
        bad = [c for c in project.condensation_results
               if c.condensation_risk or c.normative_fail]
        add_para(f"Проверено элементов: {len(project.condensation_results)}, "
                 f"с замечаниями: {len(bad)}.")
        if bad:
            rows = [["Помещение", "Категория", "U", "τ_int, °C",
                     "t_росы, °C", "Δt факт, K", "Δt норм, K", "Риск"]]
            for c in bad[:80]:
                rows.append([
                    f"{c.space_number} {c.space_name}"[:30], c.category,
                    f"{c.u_value:.2f}", f"{c.t_surface:.1f}",
                    f"{c.t_dew:.1f}", f"{c.dt_actual:.1f}",
                    f"{c.dt_normative:.1f}",
                    "конденсат" if c.condensation_risk else "норматив",
                ])
            add_table(rows)

    doc.core_properties.title = (p.project_name or "Пояснительная записка")
    doc.core_properties.comments = "HVAC Calculator"
    doc.save(path)
