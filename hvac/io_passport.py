# -*- coding: utf-8 -*-
"""Паспорта вентиляционных установок (DOCX).

Один документ — паспорт на каждую систему вентиляции проекта
(VentilationSystem), страница на систему. Структура повторяет
типовую форму паспорта вентустановки (приложение к актам по
СП 73.13330 / форма АВОК): общие сведения, расчётные данные,
вентилятор, сеть, обслуживаемые помещения; поля фактических
замеров оставлены пустыми — их заполняет наладчик.

Нагрузки берутся из hvac/ahu_load.aggregate_ahus (свежий пересчёт
при экспорте), давление вентилятора — из аэродинамического расчёта
duct_networks_detailed, если сеть с именем системы построена.
"""

from __future__ import annotations

from datetime import datetime
from typing import List, TYPE_CHECKING

if TYPE_CHECKING:
    from hvac.project import HVACProject

# Видов вентустановок (VentilationSystem.kind) — человекочитаемые имена
_KIND_NAMES = {
    "ahu": "Приточная установка (AHU)",
    "supply_fan": "Приточный вентилятор",
    "exhaust_fan": "Вытяжной вентилятор",
    "local_exhaust": "Местный отсос",
}

_TYPE_NAMES = {
    "supply": "Приточная",
    "exhaust": "Вытяжная",
    "supply_exhaust": "Приточно-вытяжная",
}

# Максимум строк помещений в таблице паспорта
_MAX_SPACE_ROWS = 25


def export_ventilation_passports(project: "HVACProject", path: str) -> int:
    """Создаёт DOCX с паспортами всех вентсистем проекта.

    Возвращает число систем; 0 систем — ValueError (нечего выгружать).
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("Не установлен python-docx. "
                           "Выполните: pip install python-docx")

    systems = getattr(project, "ventilation_systems", {}) or {}
    if not systems:
        raise ValueError("В проекте нет систем вентиляции — паспорта "
                         "выгружать не из чего. Создайте системы в панели "
                         "«Системы».")

    from hvac.ahu_load import aggregate_ahus
    loads = aggregate_ahus(project)
    nets = getattr(project, "duct_networks_detailed", {}) or {}
    p = project.params

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def add_para(text: str, *, bold: bool = False, center: bool = False,
                 size: int = 0):
        par = doc.add_paragraph()
        if center:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        return par

    def add_table(rows: List[List], header_row: bool = True):
        if not rows:
            return
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = t.cell(i, j)
                cell.text = str(val)
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.font.size = Pt(9)
                        if header_row and i == 0:
                            run.bold = True
        doc.add_paragraph()

    for idx, (name, vs) in enumerate(sorted(systems.items())):
        if idx > 0:
            doc.add_page_break()

        add_para("ПАСПОРТ", bold=True, center=True, size=18)
        add_para("вентиляционной установки", center=True, size=12)
        add_para(f"Система {name}", bold=True, center=True, size=14)
        doc.add_paragraph()

        load = loads.get(name)
        spaces = [sp for sp in project.spaces
                  if getattr(sp, "system_ventilation", "") == name]
        area = sum(sp.area_m2 for sp in spaces)

        # А. Общие сведения
        add_para("А. Общие сведения", bold=True)
        add_table([
            ["Параметр", "Значение"],
            ["Объект", p.project_name or "—"],
            ["Город", p.city or "—"],
            ["Назначение установки",
             _KIND_NAMES.get(vs.kind, vs.kind or "—")],
            ["Тип системы",
             _TYPE_NAMES.get(vs.system_type, vs.system_type or "—")],
            ["Обслуживаемых помещений",
             f"{len(spaces)} ({area:,.0f} м²)".replace(",", " ")],
            ["Примечание", vs.note or "—"],
        ])

        # Б. Расчётные данные
        add_para("Б. Расчётные данные", bold=True)
        rows = [["Параметр", "Расчёт", "Факт (наладка)"]]
        if load is not None:
            rows += [
                ["Расход притока, м³/ч", f"{load.supply_m3_h:,.0f}".replace(",", " "), ""],
                ["Расход вытяжки, м³/ч", f"{load.exhaust_m3_h:,.0f}".replace(",", " "), ""],
                ["t подачи зимой, °C", f"{load.t_supply_winter:.1f}", ""],
                ["t подачи летом, °C", f"{load.t_supply_summer:.1f}", ""],
            ]
            if load.has_recovery:
                rows += [
                    ["Рекуператор, КПД зима/лето",
                     f"{load.recovery_eff_winter * 100:.0f}% / "
                     f"{load.recovery_eff_summer * 100:.0f}%", ""],
                    ["t после рекуператора зимой, °C",
                     f"{load.t_after_recovery_winter:.1f}", ""],
                ]
            if load.q_heater_w > 0:
                rows.append(["Калорифер, кВт",
                             f"{load.q_heater_kw:.1f}", ""])
            if load.q_cooler_total_w > 0:
                rows.append(["Охладитель (полн./явная), кВт",
                             f"{load.q_cooler_total_w / 1000:.1f} / "
                             f"{load.q_cooler_sens_w / 1000:.1f}", ""])
            if load.heating_circuit:
                rows.append(["Контур теплоснабжения",
                             load.heating_circuit, ""])
            if load.cooling_circuit:
                rows.append(["Контур холодоснабжения",
                             load.cooling_circuit, ""])
        else:
            rows.append(["Нагрузки не рассчитаны "
                         "(выполните расчёт вентиляции)", "—", ""])
        add_table(rows)

        # В. Вентилятор и сеть
        add_para("В. Вентилятор и сеть", bold=True)
        net = nets.get(name)
        rows = [["Параметр", "Расчёт", "Факт (наладка)"]]
        if net is not None and net.fan_flow_m3_h > 0:
            v_max = max((e.velocity_m_s for e in net.edges.values()),
                        default=0.0)
            crit = next((b for b in net.branches
                         if b.terminal_edge_id == net.critical_branch_id),
                        None)
            rows += [
                ["Расход вентилятора, м³/ч",
                 f"{net.fan_flow_m3_h:,.0f}".replace(",", " "), ""],
                ["Полное давление, Па",
                 f"{net.fan_pressure_required_pa:,.0f}".replace(",", " "), ""],
                ["Участков сети / ветвей",
                 f"{len(net.edges)} / {len(net.branches)}", ""],
                ["Диктующая ветвь",
                 crit.terminal_name if crit else "—", ""],
                ["Макс. скорость в сети, м/с", f"{v_max:.1f}", ""],
            ]
            from hvac.fan_catalog import select_fan
            pick = select_fan(net.fan_flow_m3_h,
                              net.fan_pressure_required_pa)
            if pick is not None:
                rows.append([
                    "Рекомендуемый вентилятор (каталог)",
                    f"{pick.model.name} ({pick.model.family}), "
                    f"{pick.model.power_w:.0f} Вт", ""])
        else:
            rows.append(["Аэродинамический расчёт сети не выполнен "
                         "(вкладка «Воздуховоды»)", "—", ""])
        rows += [
            ["Марка вентилятора (факт)", "", ""],
            ["Частота вращения, об/мин", "", ""],
            ["Двигатель: тип / кВт / об/мин", "", ""],
        ]
        add_table(rows)

        # Г. Обслуживаемые помещения
        add_para("Г. Обслуживаемые помещения", bold=True)
        if spaces:
            rows = [["№", "Помещение", "Приток, м³/ч", "Вытяжка, м³/ч"]]
            for sp in spaces[:_MAX_SPACE_ROWS]:
                rows.append([
                    sp.number, sp.name,
                    f"{sp.supply_m3h:.0f}", f"{sp.exhaust_m3h:.0f}"])
            if len(spaces) > _MAX_SPACE_ROWS:
                rows.append(["…",
                             f"и ещё {len(spaces) - _MAX_SPACE_ROWS} "
                             "помещений (см. ведомость воздухообменов)",
                             "", ""])
            rows.append([
                "Σ", "",
                f"{sum(sp.supply_m3h for sp in spaces):.0f}",
                f"{sum(sp.exhaust_m3h for sp in spaces):.0f}"])
            add_table(rows)
        else:
            add_para("Помещения не привязаны к системе.", center=False)

        # Подписи
        doc.add_paragraph()
        add_para(f"Расчёт: HVAC Calculator v4 · "
                 f"{datetime.now().strftime('%d.%m.%Y')}")
        add_para("Исполнитель: ______________________   "
                 "Наладчик: ______________________")

    doc.save(path)
    return len(systems)
