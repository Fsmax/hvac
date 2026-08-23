# -*- coding: utf-8 -*-
"""Проектные документы из предварительных ответов Coordera Calc.

Модуль намеренно не выполняет HVAC-расчёты. Он читает готовые ответы
``hvac-terminal-layout`` и ``hvac-route-network``, проверяет их назначение и
статус, а затем дословно агрегирует результаты для спецификации, ведомости
воздуховодов и короткого раздела пояснительной записки.

Публичная точка входа::

    docs = DesignDocs.from_files(terminal_response, route_response)
    docs.to_xlsx("design-docs.xlsx")
    docs.to_docx("design-docs.docx")
    docs.to_json("design-docs.json")

CLI::

    python -m hvac.design_docs --terminal terminal.json --route route.json \
        --out output-directory
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from decimal import Decimal, ROUND_CEILING
from io import BytesIO
import json
import math
from pathlib import Path
import sys
from tempfile import TemporaryDirectory
from typing import Any, Iterable, Mapping, Sequence

from hvac.catalogs.shnq_ducts import SHNQ_DUCTS, duct_velocity_limits


TERMINAL_RESPONSE_KIND = "hvac-terminal-layout-response"
ROUTE_RESPONSE_KIND = "hvac-route-network-response"
PRELIMINARY_STATUS = "PRELIMINARY"

JSON_OUTPUT_NAME = "design-docs.json"
XLSX_OUTPUT_NAME = "design-docs.xlsx"
DOCX_OUTPUT_NAME = "design-docs.docx"

NO_RESERVE_NOTE = "Без монтажного запаса"
ENGINEER_REVIEW_NOTE = "нормативные значения не сверены инженером"

_DIRECTIONS = ("supply", "exhaust")
_DIRECTION_LABELS = {"supply": "приток", "exhaust": "вытяжка"}
_DIRECTION_ORDER = {value: index for index, value in enumerate(_DIRECTIONS)}

_SEGMENT_KINDS = ("trunk", "branch", "riser", "connection")
_SEGMENT_KIND_LABELS = {
    "trunk": "Ствол",
    "branch": "Ответвления",
    "riser": "Стояки",
    "connection": "Подключения",
}
_SEGMENT_KIND_ORDER = {
    value: index for index, value in enumerate(_SEGMENT_KINDS)
}


class DesignDocsError(ValueError):
    """Базовая контролируемая ошибка формирования документов."""


class DesignDocsInputError(DesignDocsError):
    """Входной JSON не соответствует ожидаемому Calc-response."""


@dataclass(frozen=True, slots=True)
class TerminalPerformance:
    """Одна уникальная пара LwA/ΔP из direction-specific pick."""

    lwa_db: float | None
    dp_pa: float | None

    def as_dict(self) -> dict[str, float | None]:
        return {"lwaDb": self.lwa_db, "dpPa": self.dp_pa}


@dataclass(frozen=True, slots=True)
class EquipmentRow:
    """Позиция спецификации оборудования по ГОСТ 21.110."""

    position: int
    model: str
    size: str
    direction: str
    name: str
    type_mark: str
    quantity: int
    unit: str
    note: str
    selected_quantity: int
    performance_modes: tuple[TerminalPerformance, ...]
    source_spaces: tuple[str, ...]
    issue_codes: tuple[str, ...]

    def as_dict(self) -> dict[str, Any]:
        return {
            "position": self.position,
            "model": self.model,
            "size": self.size,
            "direction": self.direction,
            "directionLabel": _DIRECTION_LABELS[self.direction],
            "name": self.name,
            "typeMark": self.type_mark,
            "quantity": self.quantity,
            "unit": self.unit,
            "note": self.note,
            "selectedQuantity": self.selected_quantity,
            "performanceModes": [mode.as_dict() for mode in self.performance_modes],
            "sourceSpaces": list(self.source_spaces),
            "issueCodes": list(self.issue_codes),
        }


@dataclass(frozen=True, slots=True)
class DuctRow:
    """Строка ведомости для одного сечения и одного вида участка."""

    position: int
    shape: str
    size_label: str
    diameter_mm: int | None
    width_mm: int | None
    height_mm: int | None
    segment_kind: str
    segment_count: int
    source_length_mm: float
    length_m: float
    unit: str
    note: str

    def as_dict(self) -> dict[str, Any]:
        size: dict[str, int]
        if self.shape == "round":
            size = {"dMm": int(self.diameter_mm or 0)}
        else:
            size = {
                "wMm": int(self.width_mm or 0),
                "hMm": int(self.height_mm or 0),
            }
        return {
            "position": self.position,
            "shape": self.shape,
            "size": size,
            "sizeLabel": self.size_label,
            "segmentKind": self.segment_kind,
            "segmentKindLabel": _SEGMENT_KIND_LABELS[self.segment_kind],
            "segmentCount": self.segment_count,
            "sourceLengthMm": self.source_length_mm,
            "lengthM": self.length_m,
            "unit": self.unit,
            "note": self.note,
        }


@dataclass(frozen=True, slots=True)
class DuctSizeTotal:
    """Контрольный итог всех видов участков одного сечения."""

    shape: str
    size_label: str
    diameter_mm: int | None
    width_mm: int | None
    height_mm: int | None
    source_length_mm: float
    rounded_length_m: float

    def as_dict(self) -> dict[str, Any]:
        size: dict[str, int]
        if self.shape == "round":
            size = {"dMm": int(self.diameter_mm or 0)}
        else:
            size = {
                "wMm": int(self.width_mm or 0),
                "hMm": int(self.height_mm or 0),
            }
        return {
            "shape": self.shape,
            "size": size,
            "sizeLabel": self.size_label,
            "sourceLengthMm": self.source_length_mm,
            "roundedLengthM": self.rounded_length_m,
        }


@dataclass(frozen=True, slots=True)
class SystemDesignData:
    """Готовые расчётные данные одной системы для записки."""

    system_id: str
    airflow_m3h: float
    fan_pressure_pa: float
    critical_path_segment_ids: tuple[str, ...]
    critical_path_segment_count: int
    critical_path_length_m: float
    terminal_count: int

    def as_dict(self) -> dict[str, Any]:
        return {
            "systemId": self.system_id,
            "airflowM3h": self.airflow_m3h,
            "fanPressurePa": self.fan_pressure_pa,
            "dictatingBranch": {
                "segmentIds": list(self.critical_path_segment_ids),
                "segmentCount": self.critical_path_segment_count,
                "lengthM": self.critical_path_length_m,
            },
            "terminalCount": self.terminal_count,
        }


@dataclass(frozen=True, slots=True)
class NormReference:
    """Ссылка на применимую запись каталога скоростей."""

    key: str
    value: float | None
    minimum: float | None
    maximum: float | None
    unit: str
    status: str
    source_document: str
    source_edition: str
    source_clause: str
    source_page_pdf: int
    source_table: str | None
    note_ru: str

    def as_dict(self) -> dict[str, Any]:
        return {
            "key": self.key,
            "value": self.value,
            "range": (
                None
                if self.minimum is None and self.maximum is None
                else {"min": self.minimum, "max": self.maximum}
            ),
            "unit": self.unit,
            "status": self.status,
            "source": {
                "document": self.source_document,
                "edition": self.source_edition,
                "clause": self.source_clause,
                "pagePdf": self.source_page_pdf,
                "table": self.source_table,
            },
            "noteRu": self.note_ru,
        }


@dataclass(frozen=True, slots=True)
class VelocityNormLookup:
    """Результат вызова duct_velocity_limits для вида участка."""

    kind: str
    building: str
    observed_velocity_min_ms: float
    observed_velocity_max_ms: float
    references: tuple[NormReference, ...]

    def as_dict(self) -> dict[str, Any]:
        catalog_statuses = sorted(
            {
                entry.status
                for entry in SHNQ_DUCTS.entries
                if entry.key.startswith("velocity.")
            }
        )
        return {
            "kind": self.kind,
            "building": self.building,
            "observedVelocityMs": {
                "min": self.observed_velocity_min_ms,
                "max": self.observed_velocity_max_ms,
            },
            "lookupStatus": "FOUND" if self.references else "NO_APPLICABLE_VALUE",
            "catalogSource": {
                "document": SHNQ_DUCTS.document,
                "edition": SHNQ_DUCTS.edition,
            },
            "catalogStatuses": catalog_statuses,
            "references": [reference.as_dict() for reference in self.references],
        }


def _reject_constant(value: str) -> None:
    raise DesignDocsInputError(f"Недопустимая JSON-константа: {value}")


def _object_without_duplicates(pairs: list[tuple[str, Any]]) -> dict[str, Any]:
    result: dict[str, Any] = {}
    for key, value in pairs:
        if key in result:
            raise DesignDocsInputError(f"Повторяющийся JSON-ключ: {key}")
        result[key] = value
    return result


def _load_response(
    path: str | Path,
    *,
    expected_kind: str,
    label: str,
) -> dict[str, Any]:
    source = Path(path)
    try:
        raw = source.read_text(encoding="utf-8")
    except (OSError, UnicodeError) as exc:
        raise DesignDocsInputError(f"Не удалось прочитать {label}: {source}") from exc

    try:
        value = json.loads(
            raw,
            object_pairs_hook=_object_without_duplicates,
            parse_constant=_reject_constant,
        )
    except (json.JSONDecodeError, UnicodeError) as exc:
        raise DesignDocsInputError(f"Некорректный JSON {label}: {source}: {exc}") from exc

    if not isinstance(value, dict):
        raise DesignDocsInputError(f"{label}: ожидался JSON-объект верхнего уровня")
    if value.get("kind") != expected_kind:
        raise DesignDocsInputError(
            f"{label}: ожидался kind={expected_kind!r}, получен {value.get('kind')!r}"
        )
    if value.get("status") != PRELIMINARY_STATUS:
        raise DesignDocsInputError(
            f"{label}: ожидался status={PRELIMINARY_STATUS!r}, "
            f"получен {value.get('status')!r}"
        )
    return value


def _as_object(value: Any, location: str) -> dict[str, Any]:
    if not isinstance(value, dict):
        raise DesignDocsInputError(f"{location}: ожидался объект")
    return value


def _as_list(value: Any, location: str) -> list[Any]:
    if not isinstance(value, list):
        raise DesignDocsInputError(f"{location}: ожидался массив")
    return value


def _as_string(value: Any, location: str) -> str:
    if not isinstance(value, str) or not value.strip():
        raise DesignDocsInputError(f"{location}: ожидалась непустая строка")
    return value


def _as_number(value: Any, location: str) -> float:
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        raise DesignDocsInputError(f"{location}: ожидалось число")
    number = float(value)
    if not math.isfinite(number):
        raise DesignDocsInputError(f"{location}: ожидалось конечное число")
    return number


def _as_optional_number(value: Any, location: str) -> float | None:
    if value is None:
        return None
    return _as_number(value, location)


def _as_non_negative_int(value: Any, location: str) -> int:
    number = _as_number(value, location)
    if number < 0 or not number.is_integer():
        raise DesignDocsInputError(f"{location}: ожидалось целое число >= 0")
    return int(number)


def _as_positive_size(value: Any, location: str) -> int:
    size = _as_non_negative_int(value, location)
    if size <= 0:
        raise DesignDocsInputError(f"{location}: размер должен быть больше нуля")
    return size


def _optional_number_sort_key(value: float | None) -> tuple[int, float]:
    return (1, 0.0) if value is None else (0, value)


def _fmt_number(value: float | None, *, decimals: int = 2) -> str:
    if value is None:
        return "—"
    text = f"{value:.{decimals}f}".rstrip("0").rstrip(".")
    return text or "0"


def _performance_text(modes: tuple[TerminalPerformance, ...]) -> str:
    labels = []
    for mode in modes:
        labels.append(
            f"LwA={_fmt_number(mode.lwa_db)} дБА; "
            f"ΔP={_fmt_number(mode.dp_pa)} Па"
        )
    return " / ".join(labels)


def _issue_note(room_name: str, issue_codes: Iterable[str]) -> list[str]:
    notes = []
    for code in sorted(set(issue_codes)):
        if code == "TERMINAL_GRID_OVERFLOW":
            notes.append(f"{room_name}: недоразмещено, см. {code}")
        else:
            notes.append(f"{room_name}: см. {code}")
    return notes


def _build_specification(terminal: Mapping[str, Any]) -> tuple[EquipmentRow, ...]:
    results = _as_list(terminal.get("results"), "terminal.results")
    groups: dict[tuple[str, str, str], dict[str, Any]] = {}

    for result_index, raw_result in enumerate(results):
        location = f"terminal.results[{result_index}]"
        result = _as_object(raw_result, location)
        space_id = _as_string(result.get("spaceId"), f"{location}.spaceId")
        room_name = _as_string(result.get("name"), f"{location}.name")

        placements_by_direction = {direction: 0 for direction in _DIRECTIONS}
        for placement_index, raw_placement in enumerate(
            _as_list(result.get("placements"), f"{location}.placements")
        ):
            placement = _as_object(
                raw_placement, f"{location}.placements[{placement_index}]"
            )
            direction = placement.get("direction")
            if direction not in placements_by_direction:
                raise DesignDocsInputError(
                    f"{location}.placements[{placement_index}].direction: "
                    f"неизвестное направление {direction!r}"
                )
            placements_by_direction[direction] += 1

        issue_codes: list[str] = []
        for issue_index, raw_issue in enumerate(
            _as_list(result.get("issues"), f"{location}.issues")
        ):
            issue = _as_object(raw_issue, f"{location}.issues[{issue_index}]")
            issue_codes.append(
                _as_string(issue.get("code"), f"{location}.issues[{issue_index}].code")
            )

        for direction in _DIRECTIONS:
            pick_key = "supplyPick" if direction == "supply" else "exhaustPick"
            raw_pick = result.get(pick_key)
            if raw_pick is None:
                if placements_by_direction[direction]:
                    raise DesignDocsInputError(
                        f"{location}: есть placements={direction}, но {pick_key}=null"
                    )
                continue

            pick = _as_object(raw_pick, f"{location}.{pick_key}")
            model = _as_string(pick.get("model"), f"{location}.{pick_key}.model")
            size = _as_string(pick.get("size"), f"{location}.{pick_key}.size")
            selected = _as_non_negative_int(
                pick.get("nUnits"), f"{location}.{pick_key}.nUnits"
            )
            applicable_issue_codes = [
                code
                for code in issue_codes
                if code != "TERMINAL_GRID_OVERFLOW"
                or placements_by_direction[direction] < selected
            ]
            mode = TerminalPerformance(
                lwa_db=_as_optional_number(
                    pick.get("lwaDb"), f"{location}.{pick_key}.lwaDb"
                ),
                dp_pa=_as_optional_number(
                    pick.get("dpPa"), f"{location}.{pick_key}.dpPa"
                ),
            )

            key = (model, size, direction)
            group = groups.setdefault(
                key,
                {
                    "quantity": 0,
                    "selected": 0,
                    "modes": set(),
                    "spaces": set(),
                    "issue_codes": set(),
                    "notes": set(),
                },
            )
            group["quantity"] += placements_by_direction[direction]
            group["selected"] += selected
            group["modes"].add(mode)
            group["spaces"].add(space_id)
            group["issue_codes"].update(applicable_issue_codes)
            group["notes"].update(
                _issue_note(room_name, applicable_issue_codes)
            )

    ordered_keys = sorted(
        groups,
        key=lambda key: (key[0].casefold(), key[1].casefold(), _DIRECTION_ORDER[key[2]]),
    )
    rows = []
    for position, key in enumerate(ordered_keys, 1):
        model, size, direction = key
        group = groups[key]
        modes = tuple(
            sorted(
                group["modes"],
                key=lambda mode: (
                    _optional_number_sort_key(mode.lwa_db),
                    _optional_number_sort_key(mode.dp_pa),
                ),
            )
        )
        name = (
            f"Воздухораспределитель {model}, типоразмер {size}; "
            f"{_performance_text(modes)}"
        )
        rows.append(
            EquipmentRow(
                position=position,
                model=model,
                size=size,
                direction=direction,
                name=name,
                type_mark=f"{model} ({_DIRECTION_LABELS[direction]})",
                quantity=group["quantity"],
                unit="шт.",
                note="; ".join(sorted(group["notes"])),
                selected_quantity=group["selected"],
                performance_modes=modes,
                source_spaces=tuple(sorted(group["spaces"])),
                issue_codes=tuple(sorted(group["issue_codes"])),
            )
        )
    return tuple(rows)


def _size_key(
    segment: Mapping[str, Any], location: str
) -> tuple[str, int, int]:
    shape = segment.get("shape")
    size = _as_object(segment.get("size"), f"{location}.size")
    if shape == "round":
        return ("round", _as_positive_size(size.get("dMm"), f"{location}.size.dMm"), 0)
    if shape == "rect":
        return (
            "rect",
            _as_positive_size(size.get("wMm"), f"{location}.size.wMm"),
            _as_positive_size(size.get("hMm"), f"{location}.size.hMm"),
        )
    raise DesignDocsInputError(f"{location}.shape: неизвестная форма {shape!r}")


def _size_label(key: tuple[str, int, int]) -> str:
    shape, first, second = key
    return f"Ø{first}" if shape == "round" else f"{first}×{second}"


def _size_sort_key(key: tuple[str, int, int]) -> tuple[int, int, int]:
    shape, first, second = key
    return (0 if shape == "round" else 1, first, second)


def _ceil_tenth_m(length_mm: Decimal) -> Decimal:
    return (length_mm / Decimal("1000")).quantize(
        Decimal("0.1"), rounding=ROUND_CEILING
    )


def _decimal_as_float(value: Decimal) -> float:
    return float(value)


def _parse_segments(
    route: Mapping[str, Any],
) -> tuple[
    tuple[DuctRow, ...],
    tuple[DuctSizeTotal, ...],
    dict[str, dict[str, Any]],
]:
    segments = _as_list(route.get("segments"), "route.segments")
    grouped: dict[tuple[tuple[str, int, int], str], dict[str, Any]] = {}
    by_size: dict[tuple[str, int, int], Decimal] = {}
    segment_by_id: dict[str, dict[str, Any]] = {}

    for segment_index, raw_segment in enumerate(segments):
        location = f"route.segments[{segment_index}]"
        segment = _as_object(raw_segment, location)
        segment_id = _as_string(segment.get("segmentId"), f"{location}.segmentId")
        if segment_id in segment_by_id:
            raise DesignDocsInputError(f"{location}.segmentId: дубликат {segment_id!r}")
        segment_by_id[segment_id] = segment

        kind = segment.get("kind")
        if kind not in _SEGMENT_KIND_ORDER:
            raise DesignDocsInputError(f"{location}.kind: неизвестный вид {kind!r}")
        key = _size_key(segment, location)
        length_number = _as_number(segment.get("lengthMm"), f"{location}.lengthMm")
        if length_number < 0:
            raise DesignDocsInputError(f"{location}.lengthMm: длина не может быть < 0")
        length_mm = Decimal(str(segment.get("lengthMm")))

        record = grouped.setdefault((key, kind), {"length_mm": Decimal(0), "count": 0})
        record["length_mm"] += length_mm
        record["count"] += 1
        by_size[key] = by_size.get(key, Decimal(0)) + length_mm

        _as_string(segment.get("systemId"), f"{location}.systemId")
        _as_string(segment.get("a"), f"{location}.a")
        _as_string(segment.get("b"), f"{location}.b")
        _as_number(segment.get("velocityMs"), f"{location}.velocityMs")

    ordered_groups = sorted(
        grouped,
        key=lambda item: (_size_sort_key(item[0]), _SEGMENT_KIND_ORDER[item[1]]),
    )
    rows = []
    for position, (key, kind) in enumerate(ordered_groups, 1):
        shape, first, second = key
        record = grouped[(key, kind)]
        rows.append(
            DuctRow(
                position=position,
                shape=shape,
                size_label=_size_label(key),
                diameter_mm=first if shape == "round" else None,
                width_mm=first if shape == "rect" else None,
                height_mm=second if shape == "rect" else None,
                segment_kind=kind,
                segment_count=record["count"],
                source_length_mm=_decimal_as_float(record["length_mm"]),
                length_m=_decimal_as_float(_ceil_tenth_m(record["length_mm"])),
                unit="м",
                note=NO_RESERVE_NOTE,
            )
        )

    totals = []
    for key in sorted(by_size, key=_size_sort_key):
        shape, first, second = key
        scheduled_length_m = sum(
            (
                _ceil_tenth_m(record["length_mm"])
                for (group_key, _kind), record in grouped.items()
                if group_key == key
            ),
            Decimal(0),
        )
        totals.append(
            DuctSizeTotal(
                shape=shape,
                size_label=_size_label(key),
                diameter_mm=first if shape == "round" else None,
                width_mm=first if shape == "rect" else None,
                height_mm=second if shape == "rect" else None,
                source_length_mm=_decimal_as_float(by_size[key]),
                rounded_length_m=_decimal_as_float(scheduled_length_m),
            )
        )
    return tuple(rows), tuple(totals), segment_by_id


def _build_systems(
    route: Mapping[str, Any],
    segment_by_id: Mapping[str, Mapping[str, Any]],
) -> tuple[SystemDesignData, ...]:
    node_kinds: dict[str, str] = {}
    for node_index, raw_node in enumerate(_as_list(route.get("nodes"), "route.nodes")):
        location = f"route.nodes[{node_index}]"
        node = _as_object(raw_node, location)
        node_id = _as_string(node.get("nodeId"), f"{location}.nodeId")
        if node_id in node_kinds:
            raise DesignDocsInputError(f"{location}.nodeId: дубликат {node_id!r}")
        node_kinds[node_id] = _as_string(node.get("kind"), f"{location}.kind")

    terminal_nodes = {node_id for node_id, kind in node_kinds.items() if kind == "terminal"}
    systems = []
    seen_system_ids: set[str] = set()
    for system_index, raw_system in enumerate(
        _as_list(route.get("systems"), "route.systems")
    ):
        location = f"route.systems[{system_index}]"
        system = _as_object(raw_system, location)
        system_id = _as_string(system.get("systemId"), f"{location}.systemId")
        if system_id in seen_system_ids:
            raise DesignDocsInputError(f"{location}.systemId: дубликат {system_id!r}")
        seen_system_ids.add(system_id)

        critical_ids = tuple(
            _as_string(value, f"{location}.criticalPathSegmentIds[{index}]")
            for index, value in enumerate(
                _as_list(
                    system.get("criticalPathSegmentIds"),
                    f"{location}.criticalPathSegmentIds",
                )
            )
        )
        if len(set(critical_ids)) != len(critical_ids):
            raise DesignDocsInputError(
                f"{location}.criticalPathSegmentIds: участки должны быть уникальны"
            )
        critical_length_mm = Decimal(0)
        for critical_id in critical_ids:
            segment = segment_by_id.get(critical_id)
            if segment is None:
                raise DesignDocsInputError(
                    f"{location}: criticalPathSegmentIds содержит неизвестный "
                    f"участок {critical_id!r}"
                )
            if segment.get("systemId") != system_id:
                raise DesignDocsInputError(
                    f"{location}: участок {critical_id!r} принадлежит другой системе"
                )
            critical_length_mm += Decimal(str(segment["lengthMm"]))

        incident_terminal_nodes: set[str] = set()
        for segment in segment_by_id.values():
            if segment.get("systemId") != system_id:
                continue
            for endpoint in (segment.get("a"), segment.get("b")):
                if endpoint in terminal_nodes:
                    incident_terminal_nodes.add(str(endpoint))

        systems.append(
            SystemDesignData(
                system_id=system_id,
                airflow_m3h=_as_number(
                    system.get("fanFlowM3h"), f"{location}.fanFlowM3h"
                ),
                fan_pressure_pa=_as_number(
                    system.get("fanPressurePa"), f"{location}.fanPressurePa"
                ),
                critical_path_segment_ids=critical_ids,
                critical_path_segment_count=len(critical_ids),
                critical_path_length_m=_decimal_as_float(
                    critical_length_mm / Decimal("1000")
                ),
                terminal_count=len(incident_terminal_nodes),
            )
        )
    return tuple(sorted(systems, key=lambda system: system.system_id.casefold()))


def _norm_reference(entry: Any) -> NormReference:
    range_value = entry.range
    return NormReference(
        key=entry.key,
        value=entry.value,
        minimum=None if range_value is None else range_value.minimum,
        maximum=None if range_value is None else range_value.maximum,
        unit=entry.unit,
        status=entry.status,
        source_document=entry.source.document,
        source_edition=entry.source.edition,
        source_clause=entry.source.clause,
        source_page_pdf=entry.source.page_pdf,
        source_table=entry.source.table,
        note_ru=entry.note_ru,
    )


def _build_velocity_norm_lookups(
    segment_by_id: Mapping[str, Mapping[str, Any]],
) -> tuple[VelocityNormLookup, ...]:
    velocities: dict[str, list[float]] = {}
    for index, segment in enumerate(segment_by_id.values()):
        kind = str(segment.get("kind"))
        velocity = _as_number(segment.get("velocityMs"), f"route.segments[{index}].velocityMs")
        velocities.setdefault(kind, []).append(velocity)

    lookups = []
    for kind in sorted(velocities, key=lambda value: _SEGMENT_KIND_ORDER[value]):
        references = tuple(
            sorted(
                (_norm_reference(entry) for entry in duct_velocity_limits(kind, "all")),
                key=lambda reference: reference.key,
            )
        )
        lookups.append(
            VelocityNormLookup(
                kind=kind,
                building="all",
                observed_velocity_min_ms=min(velocities[kind]),
                observed_velocity_max_ms=max(velocities[kind]),
                references=references,
            )
        )
    return tuple(lookups)


def _source_summary(response: Mapping[str, Any]) -> dict[str, Any]:
    engine = response.get("engine")
    engine_id = engine.get("id") if isinstance(engine, dict) else None
    engine_version = engine.get("version") if isinstance(engine, dict) else None
    return {
        "kind": response.get("kind"),
        "status": response.get("status"),
        "requestId": response.get("requestId"),
        "responseHash": response.get("responseHash"),
        "engine": {"id": engine_id, "version": engine_version},
    }


class DesignDocs:
    """Неизменяемый набор данных и экспортов проектных документов."""

    def __init__(self, terminal_path: str | Path, route_path: str | Path) -> None:
        terminal = _load_response(
            terminal_path,
            expected_kind=TERMINAL_RESPONSE_KIND,
            label="terminal-layout response",
        )
        route = _load_response(
            route_path,
            expected_kind=ROUTE_RESPONSE_KIND,
            label="route-network response",
        )

        specification = _build_specification(terminal)
        ducts, duct_totals, segment_by_id = _parse_segments(route)
        systems = _build_systems(route, segment_by_id)
        velocity_norm_lookups = _build_velocity_norm_lookups(segment_by_id)

        self.status = PRELIMINARY_STATUS
        self.specification = specification
        self.ducts = ducts
        self.duct_totals = duct_totals
        self.systems = systems
        self.velocity_norm_lookups = velocity_norm_lookups
        self.velocity_norms = tuple(
            reference
            for lookup in velocity_norm_lookups
            for reference in lookup.references
        )
        self.normative_note = ENGINEER_REVIEW_NOTE
        self._terminal_source = _source_summary(terminal)
        self._route_source = _source_summary(route)

    @classmethod
    def from_files(
        cls, terminal_path: str | Path, route_path: str | Path
    ) -> "DesignDocs":
        """Загрузить два Calc-response и сформировать структуру документов."""

        return cls(terminal_path, route_path)

    @property
    def duct_schedule(self) -> tuple[DuctRow, ...]:
        """Совместимый читаемый алиас для строк ведомости."""

        return self.ducts

    def as_dict(self) -> dict[str, Any]:
        """Вернуть детерминированную JSON-совместимую структуру."""

        catalog_velocity_statuses = sorted(
            {
                entry.status
                for entry in SHNQ_DUCTS.entries
                if entry.key.startswith("velocity.")
            }
        )
        total_source_mm = sum(
            (Decimal(str(row.source_length_mm)) for row in self.ducts),
            Decimal(0),
        )
        total_scheduled_m = sum(
            (Decimal(str(row.length_m)) for row in self.ducts),
            Decimal(0),
        )
        return {
            "schemaVersion": "1.0",
            "kind": "hvac-design-documents",
            "status": self.status,
            "sources": {
                "terminalLayout": dict(self._terminal_source),
                "routeNetwork": dict(self._route_source),
            },
            "specification": {
                "title": "Спецификация оборудования",
                "standard": "ГОСТ 21.110-2013",
                "items": [row.as_dict() for row in self.specification],
            },
            "ductSchedule": {
                "title": "Ведомость воздуховодов",
                "rounding": "Сумма группы округлена вверх до 0,1 м",
                "reserveAdded": False,
                "note": NO_RESERVE_NOTE,
                "items": [row.as_dict() for row in self.ducts],
                "totalsBySize": [row.as_dict() for row in self.duct_totals],
                "totalSourceLengthM": _decimal_as_float(
                    total_source_mm / Decimal("1000")
                ),
                "totalScheduledLengthM": _decimal_as_float(total_scheduled_m),
            },
            "designNote": {
                "title": "Вентиляция: принятые решения",
                "systems": [system.as_dict() for system in self.systems],
                "velocityNorms": {
                    "catalog": {
                        "document": SHNQ_DUCTS.document,
                        "edition": SHNQ_DUCTS.edition,
                        "velocityEntryStatuses": catalog_velocity_statuses,
                    },
                    "buildingLookup": "all",
                    "lookups": [
                        lookup.as_dict() for lookup in self.velocity_norm_lookups
                    ],
                    "note": (
                        "Для видов участков route-response применимые численные "
                        "значения duct_velocity_limits не найдены; "
                        f"{ENGINEER_REVIEW_NOTE}."
                    ),
                },
            },
        }

    def to_json(self, path: str | Path) -> None:
        """Записать структуру в UTF-8 JSON без недетерминированных полей."""

        payload = (
            json.dumps(self.as_dict(), ensure_ascii=False, indent=2, sort_keys=True)
            + "\n"
        ).encode("utf-8")
        _write_once(Path(path), payload)

    def to_xlsx(self, path: str | Path) -> None:
        """Создать XLSX с листами «Спецификация» и «Воздуховоды»."""

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
            from openpyxl.utils import get_column_letter
        except ImportError as exc:  # pragma: no cover - dependency is declared
            raise RuntimeError("Не установлен openpyxl") from exc

        workbook = Workbook()
        specification_sheet = workbook.active
        specification_sheet.title = "Спецификация"

        thin = Side(border_style="thin", color="666666")
        border = Border(top=thin, left=thin, right=thin, bottom=thin)
        head_fill = PatternFill("solid", fgColor="1F4E78")
        head_font = Font(bold=True, color="FFFFFF")
        total_fill = PatternFill("solid", fgColor="DCE6F1")
        total_font = Font(bold=True, color="333333")

        _xlsx_title(
            specification_sheet,
            "СПЕЦИФИКАЦИЯ ОБОРУДОВАНИЯ",
            "Статус исходных расчётов: PRELIMINARY",
            "(по ГОСТ 21.110-2013)",
            6,
            Font,
            Alignment,
        )
        specification_headers = [
            "Поз.",
            "Наименование",
            "Тип, марка",
            "Кол-во",
            "Ед.",
            "Примечание",
        ]
        specification_sheet.append([])
        specification_sheet.append(specification_headers)
        specification_header_row = specification_sheet.max_row
        _style_xlsx_header(
            specification_sheet[specification_header_row],
            head_fill,
            head_font,
            border,
            Alignment,
        )
        for row in self.specification:
            specification_sheet.append(
                [
                    row.position,
                    row.name,
                    row.type_mark,
                    row.quantity,
                    row.unit,
                    row.note,
                ]
            )
            _style_xlsx_body(
                specification_sheet[specification_sheet.max_row], border, Alignment
            )
        _set_xlsx_widths(
            specification_sheet, [7, 58, 30, 10, 10, 46], get_column_letter
        )
        specification_sheet.freeze_panes = f"A{specification_header_row + 1}"
        specification_sheet.auto_filter.ref = (
            f"A{specification_header_row}:F{specification_sheet.max_row}"
        )

        duct_sheet = workbook.create_sheet("Воздуховоды")
        _xlsx_title(
            duct_sheet,
            "ВЕДОМОСТЬ ВОЗДУХОВОДОВ",
            f"{NO_RESERVE_NOTE}. Длины групп округлены вверх до 0,1 м.",
            "Исходные данные: hvac-route-network response (PRELIMINARY)",
            7,
            Font,
            Alignment,
        )
        duct_headers = [
            "Поз.",
            "Сечение",
            "Категория",
            "Участков",
            "Длина, м",
            "Ед.",
            "Примечание",
        ]
        duct_sheet.append([])
        duct_sheet.append(duct_headers)
        duct_header_row = duct_sheet.max_row
        _style_xlsx_header(
            duct_sheet[duct_header_row], head_fill, head_font, border, Alignment
        )
        for row in self.ducts:
            duct_sheet.append(
                [
                    row.position,
                    row.size_label,
                    _SEGMENT_KIND_LABELS[row.segment_kind],
                    row.segment_count,
                    row.length_m,
                    row.unit,
                    row.note,
                ]
            )
            body = duct_sheet[duct_sheet.max_row]
            _style_xlsx_body(body, border, Alignment)
            body[4].number_format = "0.0"
        for total in self.duct_totals:
            segment_count = sum(
                row.segment_count
                for row in self.ducts
                if row.shape == total.shape and row.size_label == total.size_label
            )
            duct_sheet.append(
                [
                    "",
                    total.size_label,
                    "Итого по сечению",
                    segment_count,
                    total.rounded_length_m,
                    "м",
                    NO_RESERVE_NOTE,
                ]
            )
            total_row = duct_sheet[duct_sheet.max_row]
            _style_xlsx_body(total_row, border, Alignment)
            for cell in total_row:
                cell.fill = total_fill
                cell.font = total_font
            total_row[4].number_format = "0.0"
        duct_sheet.append(
            [
                "",
                "Все сечения",
                "ИТОГО ПО ВЕДОМОСТИ",
                sum(row.segment_count for row in self.ducts),
                sum(row.length_m for row in self.ducts),
                "м",
                NO_RESERVE_NOTE,
            ]
        )
        grand_total_row = duct_sheet[duct_sheet.max_row]
        _style_xlsx_body(grand_total_row, border, Alignment)
        for cell in grand_total_row:
            cell.fill = head_fill
            cell.font = head_font
        grand_total_row[4].number_format = "0.0"
        _set_xlsx_widths(duct_sheet, [7, 16, 22, 12, 14, 10, 34], get_column_letter)
        duct_sheet.freeze_panes = f"A{duct_header_row + 1}"
        duct_sheet.auto_filter.ref = f"A{duct_header_row}:G{duct_sheet.max_row}"

        buffer = BytesIO()
        workbook.save(buffer)
        _write_once(Path(path), buffer.getvalue())

    def to_docx(self, path: str | Path) -> None:
        """Создать короткий раздел «Вентиляция: принятые решения»."""

        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as exc:  # pragma: no cover - dependency is declared
            raise RuntimeError("Не установлен python-docx") from exc

        document = Document()
        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10)

        document.add_heading("Вентиляция: принятые решения", level=1)
        document.add_paragraph(
            "Раздел сформирован по готовым ответам hvac-terminal-layout и "
            "hvac-route-network со статусом PRELIMINARY; расчётные значения "
            "в документе не пересчитывались."
        )

        document.add_heading("Расчётные данные систем", level=2)
        system_rows = [
            [
                "Система",
                "Расход, м³/ч",
                "Давление вентилятора, Па",
                "Диктующая ветвь",
                "Терминалов",
            ]
        ]
        for system in self.systems:
            system_rows.append(
                [
                    system.system_id,
                    _fmt_number(system.airflow_m3h, decimals=3),
                    _fmt_number(system.fan_pressure_pa, decimals=6),
                    (
                        f"{system.critical_path_segment_count} участков; "
                        f"{_fmt_number(system.critical_path_length_m, decimals=3)} м"
                    ),
                    str(system.terminal_count),
                ]
            )
        _add_docx_table(document, system_rows, Pt)

        document.add_heading("Воздуховоды", level=2)
        document.add_paragraph(
            f"Длины приняты буквально из сегментов трассы. {NO_RESERVE_NOTE}. "
            "Каждая группа сечения и категории округлена вверх до 0,1 м."
        )
        duct_rows = [["Сечение", "Категория", "Участков", "Длина, м"]]
        for row in self.ducts:
            duct_rows.append(
                [
                    row.size_label,
                    _SEGMENT_KIND_LABELS[row.segment_kind],
                    str(row.segment_count),
                    f"{row.length_m:.1f}",
                ]
            )
        _add_docx_table(document, duct_rows, Pt)

        document.add_heading("Ссылки на нормы скорости", level=2)
        document.add_paragraph(
            f"Каталог: {SHNQ_DUCTS.document}, редакция {SHNQ_DUCTS.edition}. "
            "Статус доступных записей каталога скорости: unverified. "
            "Тип здания отсутствует в route-response, поэтому выполнен только "
            "явный поиск building=all для фактических видов участков."
        )
        for lookup in self.velocity_norm_lookups:
            if not lookup.references:
                document.add_paragraph(
                    f"{_SEGMENT_KIND_LABELS[lookup.kind]}: применимое численное "
                    "значение duct_velocity_limits не найдено."
                )
                continue
            for reference in lookup.references:
                document.add_paragraph(
                    f"{reference.key}: статус {reference.status}; "
                    f"{reference.source_document}, п. {reference.source_clause}, "
                    f"PDF с. {reference.source_page_pdf}."
                )
        document.add_paragraph(f"Важно: {ENGINEER_REVIEW_NOTE}.")

        document.core_properties.title = "Вентиляция: принятые решения"
        document.core_properties.comments = "HVAC Calculator; PRELIMINARY"
        buffer = BytesIO()
        document.save(buffer)
        _write_once(Path(path), buffer.getvalue())


def _write_once(path: Path, payload: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        with path.open("xb") as stream:
            stream.write(payload)
    except FileExistsError:
        raise FileExistsError(f"Файл уже существует, перезапись запрещена: {path}")


def _xlsx_title(
    sheet: Any,
    title: str,
    subtitle: str,
    caption: str,
    columns: int,
    font_type: Any,
    alignment_type: Any,
) -> None:
    sheet["A1"] = title
    sheet["A1"].font = font_type(bold=True, size=14)
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=columns)
    sheet["A1"].alignment = alignment_type(horizontal="center")
    sheet["A2"] = subtitle
    sheet["A2"].font = font_type(italic=True)
    sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=columns)
    sheet["A3"] = caption
    sheet["A3"].font = font_type(size=9, italic=True)
    sheet.merge_cells(start_row=3, start_column=1, end_row=3, end_column=columns)
    sheet["A3"].alignment = alignment_type(horizontal="center")


def _style_xlsx_header(
    cells: Iterable[Any],
    fill: Any,
    font: Any,
    border: Any,
    alignment_type: Any,
) -> None:
    for cell in cells:
        cell.fill = fill
        cell.font = font
        cell.border = border
        cell.alignment = alignment_type(
            horizontal="center", vertical="center", wrap_text=True
        )


def _style_xlsx_body(
    cells: Iterable[Any], border: Any, alignment_type: Any
) -> None:
    for cell in cells:
        cell.border = border
        cell.alignment = alignment_type(wrap_text=True, vertical="top")


def _set_xlsx_widths(sheet: Any, widths: Sequence[int], get_column_letter: Any) -> None:
    for index, width in enumerate(widths, 1):
        sheet.column_dimensions[get_column_letter(index)].width = width


def _add_docx_table(document: Any, rows: Sequence[Sequence[str]], pt_type: Any) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = pt_type(9)
                    if row_index == 0:
                        run.bold = True
    document.add_paragraph()


def build_design_docs(
    terminal_path: str | Path, route_path: str | Path
) -> DesignDocs:
    """Функциональный фасад для будущих интеграций Tools."""

    return DesignDocs.from_files(terminal_path, route_path)


def _argument_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="python -m hvac.design_docs",
        description=(
            "Сформировать XLSX, DOCX и JSON из предварительных "
            "hvac-terminal-layout/hvac-route-network responses"
        ),
    )
    parser.add_argument("--terminal", required=True, help="terminal-layout response JSON")
    parser.add_argument("--route", required=True, help="route-network response JSON")
    parser.add_argument("--out", required=True, help="новый каталог для трёх документов")
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    """CLI. Каталог ``--out`` создаётся строго один раз."""

    args = _argument_parser().parse_args(argv)
    output_directory = Path(args.out)
    try:
        documents = DesignDocs.from_files(args.terminal, args.route)
        if output_directory.exists():
            raise FileExistsError(
                f"Каталог уже существует, повторный запуск запрещён: "
                f"{output_directory}"
            )
        output_directory.parent.mkdir(parents=True, exist_ok=True)
        with TemporaryDirectory(
            prefix=".design-docs-", dir=output_directory.parent
        ) as staging_name:
            staging_directory = Path(staging_name)
            documents.to_json(staging_directory / JSON_OUTPUT_NAME)
            documents.to_xlsx(staging_directory / XLSX_OUTPUT_NAME)
            documents.to_docx(staging_directory / DOCX_OUTPUT_NAME)
            staging_directory.rename(output_directory)
    except (DesignDocsError, FileExistsError, OSError, RuntimeError) as exc:
        print(f"design-docs: error: {exc}", file=sys.stderr)
        return 2

    for name in (JSON_OUTPUT_NAME, XLSX_OUTPUT_NAME, DOCX_OUTPUT_NAME):
        print(output_directory / name)
    return 0


__all__ = [
    "DOCX_OUTPUT_NAME",
    "DesignDocs",
    "DesignDocsError",
    "DesignDocsInputError",
    "DuctRow",
    "DuctSizeTotal",
    "ENGINEER_REVIEW_NOTE",
    "EquipmentRow",
    "JSON_OUTPUT_NAME",
    "NormReference",
    "SystemDesignData",
    "TerminalPerformance",
    "VelocityNormLookup",
    "XLSX_OUTPUT_NAME",
    "build_design_docs",
    "main",
]


if __name__ == "__main__":  # pragma: no cover - exercised through subprocess
    raise SystemExit(main())
